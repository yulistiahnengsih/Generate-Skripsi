from flask import Flask, request, send_file, jsonify, render_template, redirect, url_for, send_from_directory
from werkzeug.utils import secure_filename
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

app = Flask(__name__)

UPLOAD_FOLDER = 'uploads'
PROCESSED_FOLDER = 'processed'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
if not os.path.exists(PROCESSED_FOLDER):
    os.makedirs(PROCESSED_FOLDER)

def baca_template(path_template):
    return Document(path_template)


def ekstrak_judul(dokumen):
    for para in dokumen.paragraphs:
        if para.text.strip():
            return para.text.strip()
    return "Judul Tidak Ditemukan"


def ekstrak_nama(dokumen):
    nama = None
    nomor_mahasiswa = None
    paragraf_sebelum_nomor = None

    for para in dokumen.paragraphs:
        teks = para.text.strip()
        
        # Cek apakah teks berupa angka (biasanya nomor mahasiswa)
        if teks.isdigit() and len(teks) > 7:
            nomor_mahasiswa = teks
            # Nama kemungkinan ada di paragraf sebelumnya
            paragraf_sebelum_nomor = para._element.getprevious()
            if paragraf_sebelum_nomor is not None:
                nama = paragraf_sebelum_nomor.text.strip()
            break

    if nama:
        return nama
    else:
        return "Nama Tidak Ditemukan"


def ekstrak_institusi(dokumen):
    institusi = None
    
    # Loop untuk memeriksa paragraf pada halaman pertama saja
    for para in dokumen.paragraphs:
        teks = para.text.strip()
        
        # Jika menemukan kata "Universitas", ambil teks tersebut
        if teks.lower().startswith("universitas"):
            institusi = teks
            break
        
        # Jika menemukan indikator halaman baru, berhenti loop
        if "Page Break" in para._element.xml:
            break

    if institusi:
        return institusi
    else:
        return "Institusi Tidak Ditemukan"


def ekstrak_keywords(dokumen):
    keywords = None
    
    for para in dokumen.paragraphs:
        teks = para.text.strip()
        
        # Mencari pola "Keywords:" (dengan case-insensitive)
        if teks.lower().startswith("keywords:") or teks.lower().startswith("keyword:") or teks.lower().startswith("keywords :") or teks.lower().startswith("keyword :"):
            keywords = teks[len("Keywords:"):].strip()
            break

    if keywords:
        return keywords
    else:
        return "Keywords Tidak Ditemukan"


def ekstrak_kata_kunci(dokumen):
    kata_kunci = None
    
    for para in dokumen.paragraphs:
        teks = para.text.strip()
        
        # Mencari pola "Kata kunci:" (dengan case-insensitive)
        if teks.lower().startswith("kata kunci:") or teks.lower().startswith("keyword:") or teks.lower().startswith("kata kunci :") or teks.lower().startswith("keyword :"):
            kata_kunci = teks[len("Kata kunci:"):].strip()
            break

    if kata_kunci:
        return kata_kunci
    else:
        return "Kata Kunci Tidak Ditemukan"


def ekstrak_bagian(dokumen):
    bagian = {
        "Judul": ekstrak_judul(dokumen),
        "Nama Penulis": ekstrak_nama(dokumen),
        "Institusi": ekstrak_institusi(dokumen),
        "Abstrak": [],
        "Abstract": [],
        "Pendahuluan": [],
        "Metode Penelitian": [],
        "Hasil dan Pembahasan": [],
        "Kesimpulan": [],
        "Referensi": [],
        "Keywords": ekstrak_keywords(dokumen),
        "Kata Kunci": ekstrak_kata_kunci(dokumen),
    }
    
    section_flags = {
        "dalam_abstrak": False,
        "dalam_abstract": False,
        "dalam_pendahuluan": False,
        "dalam_metode_penelitian": False,
        "dalam_hasil_dan_pembahasan": False,
        "dalam_kesimpulan": False,
        "dalam_referensi": False
    }

    current_section = None
    current_list_pendahuluan = []
    current_list_metode_penelitian = []
    current_list_hasil_dan_pembahasan = []
    current_list_kesimpulan = []
    current_list_referensi = []

    #Fungsi ini digunakan untuk mengatur ulang semua flags menjadi False
    def reset_flags(flags):
        for key in flags:
            flags[key] = False

    for para in dokumen.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        #Bagian ini digunakan untuk mengabaikan bagian yang tidak perlu diambil
        if any(keyword.lower() in text.lower() for keyword in ["keywords", "keyword", "kata kunci", "ucapan terima kasih", "lampiran"]):
            reset_flags(section_flags)
            current_section = None
            continue

        #Bagian ini digunakan untuk menentukan bagian yang sedang diambil berdasarkan kata kunci yang ada di dokumen
        if "ABSTRACT" in text:
            current_section = "Abstract"
            reset_flags(section_flags)
            section_flags["dalam_abstract"] = True
        elif "ABSTRAK" in text:
            current_section = "Abstrak"
            reset_flags(section_flags)
            section_flags["dalam_abstrak"] = True        
        elif "PENDAHULUAN" in text and (para.style.name.startswith("Heading") or para.style.name.startswith("SUB")):
            current_section = "Pendahuluan"
            reset_flags(section_flags)
            section_flags["dalam_pendahuluan"] = True
        elif "METODE PENELITIAN" in text and (para.style.name.startswith("Heading") or para.style.name.startswith("SUB")):
            current_section = "Metode Penelitian"
            reset_flags(section_flags)
            section_flags["dalam_metode_penelitian"] = True
        elif "HASIL DAN DISKUSI" in text or "HASIL DAN PEMBAHASAN" in text and (para.style.name.startswith("Heading") or para.style.name.startswith("SUB")):
            current_section = "Hasil dan Pembahasan"
            reset_flags(section_flags)
            section_flags["dalam_hasil_dan_pembahasan"] = True
        elif "KESIMPULAN" in text or "PENUTUP" in text and (para.style.name.startswith("Heading") or para.style.name.startswith("SUB")):
            current_section = "Kesimpulan"
            reset_flags(section_flags)
            section_flags["dalam_kesimpulan"] = True
        elif "REFERENSI" in text or "DAFTAR PUSTAKA" in text and (para.style.name.startswith("Heading") or para.style.name.startswith("SUB")):
            current_section = "Referensi"
            reset_flags(section_flags)
            section_flags["dalam_referensi"] = True
        # Abaikan Ekstraksi Seluruh Bagian BAB
        elif "BAB" in text:
            reset_flags(section_flags)
            current_section = None
            continue
        elif para.style.name.startswith("Heading") or para.style.name.startswith("SUB"):
            continue
        elif "Tabel" in text or "Gambar" in text:
            continue
        else:
            #Jika bagian yang sedang diambil adalah bagian yang memiliki list, maka list tersebut akan diambil sebagai satu item dalam list
            if section_flags["dalam_pendahuluan"]:
                if para.style.name.startswith("List"):
                    current_list_pendahuluan.append(text)
                else:
                    if current_list_pendahuluan:
                        bagian["Pendahuluan"].append(current_list_pendahuluan)
                        current_list_pendahuluan = []
                    bagian["Pendahuluan"].append(text)
            elif section_flags["dalam_metode_penelitian"]:
                if para.style.name.startswith("List"):
                    current_list_metode_penelitian.append(text)
                else:
                    if current_list_metode_penelitian:
                        bagian["Metode Penelitian"].append(current_list_metode_penelitian)
                        current_list_metode_penelitian = []
                    bagian["Metode Penelitian"].append(text)
            elif section_flags["dalam_hasil_dan_pembahasan"]:
                if para.style.name.startswith("List"):
                    current_list_hasil_dan_pembahasan.append(text)
                else:
                    if current_list_hasil_dan_pembahasan:
                        bagian["Hasil dan Pembahasan"].append(current_list_hasil_dan_pembahasan)
                        current_list_hasil_dan_pembahasan = []
                    bagian["Hasil dan Pembahasan"].append(text)
            elif section_flags["dalam_kesimpulan"]:
                if para.style.name.startswith("List"):
                    current_list_kesimpulan.append(text)
                else:
                    if current_list_kesimpulan:
                        bagian["Kesimpulan"].append(current_list_kesimpulan)
                        current_list_kesimpulan = []
                    bagian["Kesimpulan"].append(text)
            elif section_flags["dalam_referensi"]:
                if para.style.name.startswith("List"):
                    current_list_referensi.append(text)
                else:
                    if current_list_referensi:
                        bagian["Referensi"].append(current_list_referensi)
                        current_list_referensi = []
                    bagian["Referensi"].append(text)

            else:
                if current_section:
                    bagian[current_section].append(text)

    #Bagian ini digunakan untuk menambahkan list yang belum terakomodasi ke dalam bagian yang sesuai sebelum dijadikan satu string
    if current_list_pendahuluan:
        bagian["Pendahuluan"].append(current_list_pendahuluan)
    if current_list_metode_penelitian:
        bagian["Metode Penelitian"].append(current_list_metode_penelitian)
    if current_list_hasil_dan_pembahasan:
        bagian["Hasil dan Pembahasan"].append(current_list_hasil_dan_pembahasan)
    if current_list_kesimpulan:
        bagian["Kesimpulan"].append(current_list_kesimpulan)
    if current_list_referensi:
        bagian["Referensi"].append(current_list_referensi)

    #Memeriksa apakah ada bagian yang kosong
    #for key, value in bagian.items():
    #    if not value:
    #        return f"Bagian {key} tidak ditemukan."

    return bagian

#Fungsi ini digunakan untuk mengganti bagian yang ada di template dan menyesuaikan dengan format jurnal
def sesuaikan_dengan_template(dokumen_template, bagian):
    for para in dokumen_template.paragraphs:

        #Bagian ini digunakan untuk mengganti bagian yang ada di template sesuai dengan bagian yang ada di dokumen dengan menggunakan kata kunci yang ada di dokumen, di sertai format penulisannya sesuai dengan format jurnal
        if "JUD1" in para.text:
            para.clear()
            judul = bagian.get('Judul', 'Judul Tidak Ditemukan').upper() 
            if len(judul.split()) > 15: 
                judul = ' '.join(judul.split()[:15])
            p_judul = para.insert_paragraph_before(judul)
            run_judul = p_judul.runs[0]
            run_judul.font.size = Pt(14)
            run_judul.font.name = 'Arial'
            run_judul.bold = True
            p_judul.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        elif "NP" in para.text:
            para.clear()
            penulis = bagian.get('Nama Penulis', [])
            penulis_str = ''.join(penulis)

            run_penulis = para.add_run(penulis_str)
            run_penulis.font.size = Pt(10)
            run_penulis.font.name = 'Arial'
            run_penulis.bold = True
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        elif "INS" in para.text:
            para.clear()
            institusi = bagian.get('Institusi', [])
            institusi_str = ''.join(institusi)
            
            run_institusi = para.add_run(institusi_str)
            run_institusi.font.size = Pt(10)
            run_institusi.font.name = 'Arial'
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        elif "AA1" in para.text:
            para.clear()
            for paragraph in bagian.get('Abstract', ['Tidak ada abstrak ditemukan.']):
                p_abstract = para.insert_paragraph_before(paragraph)
                run = p_abstract.runs[0]
                run.font.size = Pt(10)
                run.font.name = 'Arial'
                run.italic = True
                p_abstract.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                p_abstract.paragraph_format.space_before = Pt(3)
                p_abstract.paragraph_format.space_after = Pt(3)
                p_abstract.paragraph_format.line_spacing = 1.15
                p_abstract.paragraph_format.first_line_indent = Inches(0.5)

        elif "AA2" in para.text:
            para.clear()
            for paragraph in bagian.get('Abstrak', ['Tidak ada abstrak ditemukan.']):
                p_abstrak = para.insert_paragraph_before(paragraph)
                p_abstrak.style.font.size = Pt(10)
                p_abstrak.style.font.name = 'Arial'
                p_abstrak.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                p_abstrak.paragraph_format.space_before = Pt(3)
                p_abstrak.paragraph_format.space_after = Pt(3)
                p_abstrak.paragraph_format.line_spacing = 1.15
                p_abstrak.paragraph_format.first_line_indent = Inches(0.5)
        
        elif "AA3" in para.text:
            para.clear()
            for item_pendahuluan in bagian.get('Pendahuluan', ['Tidak ada pendahuluan ditemukan.']):
                if isinstance(item_pendahuluan, list):
                    for i, list_item in enumerate(item_pendahuluan, 1):
                        p_pendahuluan = para.insert_paragraph_before(f"{i}. {list_item}")
                        p_pendahuluan.style.font.size = Pt(10)
                        p_pendahuluan.style.font.name = 'Arial'
                        p_pendahuluan.paragraph_format.left_indent = Inches(0.5)
                        p_pendahuluan.paragraph_format.first_line_indent = Inches(-0.15)
                        p_pendahuluan.paragraph_format.space_before = Pt(3)
                        p_pendahuluan.paragraph_format.space_after = Pt(3)
                        p_pendahuluan.paragraph_format.line_spacing = 1.15
                        p_pendahuluan.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                else:
                    p_pendahuluan = para.insert_paragraph_before(item_pendahuluan)
                    p_pendahuluan.style.font.size = Pt(10)
                    p_pendahuluan.style.font.name = 'Arial'
                    p_pendahuluan.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    p_pendahuluan.paragraph_format.space_before = Pt(3)
                    p_pendahuluan.paragraph_format.space_after = Pt(3)
                    p_pendahuluan.paragraph_format.line_spacing = 1.15
                    p_pendahuluan.paragraph_format.first_line_indent = Inches(0.5)

        elif "AA4" in para.text:
            para.clear()
            for item_metode_penelitian in bagian.get('Metode Penelitian', ['Tidak ada metode penelitian ditemukan.']):
                if isinstance(item_metode_penelitian, list):
                    for i, list_item in enumerate(item_metode_penelitian, 1):
                        p_metode_penelitian = para.insert_paragraph_before(f"{i}. {list_item}")
                        p_metode_penelitian.style.font.size = Pt(10)
                        p_metode_penelitian.style.font.name = 'Arial'
                        p_metode_penelitian.paragraph_format.left_indent = Inches(0.5)
                        p_metode_penelitian.paragraph_format.first_line_indent = Inches(-0.15)
                        p_metode_penelitian.paragraph_format.space_before = Pt(3)
                        p_metode_penelitian.paragraph_format.space_after = Pt(3)
                        p_metode_penelitian.paragraph_format.line_spacing = 1.15
                        p_metode_penelitian.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                else:
                    p_metode_penelitian = para.insert_paragraph_before(item_metode_penelitian)
                    p_metode_penelitian.style.font.size = Pt(10)
                    p_metode_penelitian.style.font.name = 'Arial'
                    p_metode_penelitian.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    p_metode_penelitian.paragraph_format.space_before = Pt(3)
                    p_metode_penelitian.paragraph_format.space_after = Pt(3)
                    p_metode_penelitian.paragraph_format.line_spacing = 1.15
                    p_metode_penelitian.paragraph_format.first_line_indent = Inches(0.5)

        elif "AA5" in para.text:
            para.clear()
            for item_hasil in bagian.get('Hasil dan Pembahasan', ['Tidak ada hasil dan pembahasan ditemukan.']):
                if isinstance(item_hasil, list):
                    for i, list_item in enumerate(item_hasil, 1):
                        p_hasil = para.insert_paragraph_before(f"{i}. {list_item}")
                        p_hasil.style.font.size = Pt(10)
                        p_hasil.style.font.name = 'Arial'
                        p_hasil.paragraph_format.left_indent = Inches(0.5)
                        p_hasil.paragraph_format.first_line_indent = Inches(-0.15)
                        p_hasil.paragraph_format.space_before = Pt(3)
                        p_hasil.paragraph_format.space_after = Pt(3)
                        p_hasil.paragraph_format.line_spacing = 1.15
                        p_hasil.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                else:
                    p_hasil = para.insert_paragraph_before(item_hasil)
                    p_hasil.style.font.size = Pt(10)
                    p_hasil.style.font.name = 'Arial'
                    p_hasil.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    p_hasil.paragraph_format.space_before = Pt(3)
                    p_hasil.paragraph_format.space_after = Pt(3)
                    p_hasil.paragraph_format.line_spacing = 1.15
                    p_hasil.paragraph_format.first_line_indent = Inches(0.5)

        elif "AA6" in para.text:
            para.clear()
            for item_kesimpulan in bagian.get('Kesimpulan', ['Tidak ada kesimpulan ditemukan.']):
                if isinstance(item_kesimpulan, list):
                    for i, list_item in enumerate(item_kesimpulan, 1):
                        p_kesimpulan = para.insert_paragraph_before(f"{i}. {list_item}")
                        p_kesimpulan.style.font.size = Pt(10)
                        p_kesimpulan.style.font.name = 'Arial'
                        p_kesimpulan.paragraph_format.left_indent = Inches(0.5)
                        p_kesimpulan.paragraph_format.first_line_indent = Inches(-0.15)
                        p_kesimpulan.paragraph_format.space_before = Pt(3)
                        p_kesimpulan.paragraph_format.space_after = Pt(3)
                        p_kesimpulan.paragraph_format.line_spacing = 1.15
                        p_kesimpulan.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                else:
                    p_kesimpulan = para.insert_paragraph_before(item_kesimpulan)
                    p_kesimpulan.style.font.size = Pt(10)
                    p_kesimpulan.style.font.name = 'Arial'
                    p_kesimpulan.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    p_kesimpulan.paragraph_format.space_before = Pt(3)
                    p_kesimpulan.paragraph_format.space_after = Pt(3)
                    p_kesimpulan.paragraph_format.line_spacing = 1.15
                    p_kesimpulan.paragraph_format.first_line_indent = Inches(0.5)

        elif "AA7" in para.text:
            para.clear()
            referensi_paragraf = bagian.get('Referensi', ['Tidak ada referensi ditemukan.'])
            
            # Inisiasi counter untuk penomoran IEEE
            nomor_referensi = 1
            
            # Iterasi setiap paragraf dalam daftar referensi
            for item_referensi in referensi_paragraf:
                if isinstance(item_referensi, list):
                    for sub_item in item_referensi:
                        # Buat paragraf baru dengan penomoran IEEE [1], [2], dst.
                        p_referensi = para.insert_paragraph_before(f"[{nomor_referensi}] {sub_item}")
                        p_referensi.style.font.size = Pt(10)
                        p_referensi.style.font.name = 'Arial'
                        p_referensi.paragraph_format.left_indent = Inches(0.22)
                        p_referensi.paragraph_format.first_line_indent = Inches(-0.22)
                        p_referensi.paragraph_format.space_before = Pt(3)
                        p_referensi.paragraph_format.space_after = Pt(3)
                        p_referensi.paragraph_format.line_spacing = 1.15
                        p_referensi.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        
                        # Increment nomor referensi
                        nomor_referensi += 1
                else:
                    # Sama seperti di atas, namun untuk item tunggal
                    p_referensi = para.insert_paragraph_before(f"[{nomor_referensi}] {item_referensi}")
                    p_referensi.style.font.size = Pt(10)
                    p_referensi.style.font.name = 'Arial'
                    p_referensi.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    p_referensi.paragraph_format.left_indent = Inches(0.22)
                    p_referensi.paragraph_format.first_line_indent = Inches(-0.22)
                    p_referensi.paragraph_format.space_before = Pt(3)
                    p_referensi.paragraph_format.space_after = Pt(3)
                    p_referensi.paragraph_format.line_spacing = 1.15
                    
                    # Increment nomor referensi
                    nomor_referensi += 1

        elif "KK1" in para.text:
            para.clear()
            run = para.add_run(f"Keywords: {bagian.get('Keywords', 'Tidak ada keywords ditemukan.')}")
            run.font.size = Pt(10)
            run.font.name = 'Arial'
            run.italic = True
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        elif "KK2" in para.text:
            para.clear()
            run = para.add_run(f"Kata Kunci: {bagian.get('Kata Kunci', 'Tidak ada kata kunci ditemukan.')}")
            run.font.size = Pt(10)
            run.font.name = 'Arial'
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def simpan_dokumen_baru(dokumen_template, path_baru):
    dokumen_template.save(path_baru)


def konversi_skripsi_ke_jurnal(path_skripsi, path_template, path_output):
    dokumen_skripsi = Document(path_skripsi)
    dokumen_template = baca_template(path_template)
    bagian = ekstrak_bagian(dokumen_skripsi)
    sesuaikan_dengan_template(dokumen_template, bagian)
    simpan_dokumen_baru(dokumen_template, path_output)


@app.route('/')
def index():
    return render_template('index.html')

@app.route('/template-ainet')
def template_ainet():
    return render_template('template_ainet.html')

@app.route('/proses-generate')
def proses_generate():
    return render_template('proses_generate.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400
    
    if file:
        filename = secure_filename(file.filename)
        skripsi_path = os.path.join(UPLOAD_FOLDER, filename)
        file.save(skripsi_path)

        template_path = "template_jurnal/Template.docx"
        output_path = os.path.join(PROCESSED_FOLDER, f"processed_{filename}")
        
        konversi_skripsi_ke_jurnal(skripsi_path, template_path, output_path)
        
        return jsonify({"message": "File uploaded and processed successfully", "download_url": f"/download/{filename}"}), 200

#Fungsi ini digunakan untuk mendownload file jurnal yang telah dihasilkan
@app.route('/download/<filename>', methods=['GET'])
def download_file(filename):
    processed_path = os.path.join(PROCESSED_FOLDER, f"processed_{filename}")
    if os.path.exists(processed_path):
        return send_file(processed_path, as_attachment=True)
    else:
        return jsonify({"error": "File not found"}), 404

if __name__ == '__main__':
    app.run(debug=True)