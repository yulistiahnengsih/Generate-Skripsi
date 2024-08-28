import unittest
from docx import Document
from io import BytesIO
import logging
import os

# Mengimpor fungsi yang akan diuji
from app import baca_template, ekstrak_judul, ekstrak_nama, ekstrak_institusi, ekstrak_keywords, ekstrak_kata_kunci, ekstrak_bagian, app, sesuaikan_dengan_template

# Mengkonfigurasi logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

class TestDocumentFunctions(unittest.TestCase):

    def test_baca_template(self):
        logging.info("Memulai test_baca_template...")
        # Buat dokumen Word dummy di memori
        doc = Document()
        doc.add_paragraph("Test Document.")
        fake_file = BytesIO()
        doc.save(fake_file)
        fake_file.seek(0)

        # Uji fungsi baca_template
        result = baca_template(fake_file)
        self.assertEqual(result.paragraphs[0].text, "Test Document.")
        logging.info("test_baca_template berhasil.\n")

    def test_ekstrak_judul_with_title(self):
        logging.info("Memulai test_ekstrak_judul_with_title...")
        # Buat dokumen Word dummy dengan judul
        doc = Document()
        doc.add_paragraph("Title: Document Title")
        fake_file = BytesIO()
        doc.save(fake_file)
        fake_file.seek(0)

        # Uji fungsi ekstrak_judul
        result = ekstrak_judul(doc)
        self.assertEqual(result, "Title: Document Title")
        logging.info("test_ekstrak_judul_with_title berhasil.\n")

    def test_ekstrak_judul_no_title(self):
        logging.info("Memulai test_ekstrak_judul_no_title...")
        # Buat dokumen Word dummy tanpa judul
        doc = Document()
        fake_file = BytesIO()
        doc.save(fake_file)
        fake_file.seek(0)

        # Uji fungsi ekstrak_judul
        result = ekstrak_judul(doc)
        self.assertEqual(result, "Judul Tidak Ditemukan")
        logging.info("test_ekstrak_judul_no_title berhasil.\n")

    def test_ekstrak_nama_dengan_nomor_mahasiswa(self):
        logging.info("Memulai test_ekstrak_nama_dengan_nomor_mahasiswa...")
        # Membuat dokumen Word dummy
        doc = Document()
        doc.add_paragraph("John Doe")  # Nama mahasiswa
        doc.add_paragraph("1234567890")  # Nomor mahasiswa
        
        # Uji ekstrak_nama
        result = ekstrak_nama(doc) # ini fungsi ini akan mengembalikan nama mahasiswa
        self.assertEqual(result, "John Doe")
        logging.info("test_ekstrak_nama_dengan_nomor_mahasiswa berhasil.\n")

    def test_ekstrak_nama_tanpa_nomor_mahasiswa(self):
        logging.info("Memulai test_ekstrak_nama_tanpa_nomor_mahasiswa...")
        # Membuat dokumen Word dummy tanpa nomor mahasiswa
        doc = Document()
        doc.add_paragraph("John Doe")  # Nama tanpa nomor mahasiswa
        
        # Uji ekstrak_nama
        result = ekstrak_nama(doc) # ini fungsi ini akan mengembalikan nama mahasiswa
        self.assertEqual(result, "Nama Tidak Ditemukan")
        logging.info("test_ekstrak_nama_tanpa_nomor_mahasiswa berhasil.\n")

    def test_ekstrak_institusi_dengan_nama_universitas(self):
        logging.info("Memulai test_ekstrak_institusi_dengan_nama_universitas...")
        # Membuat dokumen Word dummy dengan institusi
        doc = Document()
        doc.add_paragraph("Universitas Indonesia")
        
        # Uji ekstrak_institusi
        result = ekstrak_institusi(doc) # fungsi ini akan mengembalikan nama institusi
        self.assertEqual(result, "Universitas Indonesia")
        logging.info("test_ekstrak_institusi_dengan_nama_universitas berhasil.\n")

    def test_ekstrak_institusi_tanpa_nama_universitas(self):
        logging.info("Memulai test_ekstrak_institusi_tanpa_nama_universitas...")
        # Membuat dokumen Word dummy tanpa nama universitas
        doc = Document()
        doc.add_paragraph("Sekolah Tinggi Teknologi")
        
        # Uji ekstrak_institusi
        result = ekstrak_institusi(doc)
        self.assertEqual(result, "Institusi Tidak Ditemukan")
        logging.info("test_ekstrak_institusi_tanpa_nama_universitas berhasil.\n")

    def test_ekstrak_keywords_dengan_keywords(self):
        logging.info("Memulai test_ekstrak_keywords_dengan_keywords...")
        # Membuat dokumen Word dummy dengan keywords
        doc = Document()
        doc.add_paragraph("Keywords: AI, Machine Learning, Deep Learning")
        
        # Uji ekstrak_keywords
        result = ekstrak_keywords(doc)
        self.assertEqual(result, "AI, Machine Learning, Deep Learning")
        logging.info("test_ekstrak_keywords_dengan_keywords berhasil.\n")

    def test_ekstrak_keywords_tanpa_keywords(self):
        logging.info("Memulai test_ekstrak_keywords_tanpa_keywords...")
        # Membuat dokumen Word dummy tanpa keywords
        doc = Document()
        doc.add_paragraph("No keywords present in this document.")
        
        # Uji ekstrak_keywords
        result = ekstrak_keywords(doc)
        self.assertEqual(result, "Keywords Tidak Ditemukan")
        logging.info("test_ekstrak_keywords_tanpa_keywords berhasil.\n")

    def test_ekstrak_kata_kunci_dengan_kata_kunci(self):
        logging.info("Memulai test_ekstrak_kata_kunci_dengan_kata_kunci...")
        # Membuat dokumen Word dummy dengan kata kunci
        doc = Document()
        doc.add_paragraph("Kata kunci: AI, Pembelajaran Mesin, Pembelajaran Mendalam")
        
        # Uji ekstrak_kata_kunci
        result = ekstrak_kata_kunci(doc)
        self.assertEqual(result, "AI, Pembelajaran Mesin, Pembelajaran Mendalam")
        logging.info("test_ekstrak_kata_kunci_dengan_kata_kunci berhasil.\n")

    def test_ekstrak_kata_kunci_tanpa_kata_kunci(self):
        logging.info("Memulai test_ekstrak_kata_kunci_tanpa_kata_kunci...")
        # Membuat dokumen Word dummy tanpa kata kunci
        doc = Document()
        doc.add_paragraph("Tidak ada kata kunci dalam dokumen ini.")
        
        # Uji ekstrak_kata_kunci
        result = ekstrak_kata_kunci(doc)
        self.assertEqual(result, "Kata Kunci Tidak Ditemukan") # Ini fungsinya adalah untuk menguji apakah kata kunci ditemukan atau tidak
        logging.info("test_ekstrak_kata_kunci_tanpa_kata_kunci berhasil.\n")

    def test_ekstrak_bagian_not_empty(self):
        logging.info("Memulai test_ekstrak_bagian_not_empty...")
        doc = Document("data/test/sample_test.docx")

        hasil = ekstrak_bagian(doc)

        # Daftar bagian yang akan diuji
        bagian_keys = [
            "Abstrak",
            "Abstract",
            "Pendahuluan",
            "Metode Penelitian",
            "Hasil dan Pembahasan",
            "Kesimpulan",
            "Referensi"
        ]

        # Periksa setiap bagian, pastikan tidak kosong
        for key in bagian_keys:
            with self.subTest(bagian=key):
                self.assertTrue(hasil[key], f"Bagian {key} seharusnya tidak kosong.")
        logging.info("test_ekstrak_bagian_not_empty berhasil.\n")

    def test_ekstrak_bagian_empty(self):
        logging.info("Memulai test_ekstrak_bagian_empty...")
        doc = Document("data/test/sample_test_empty.docx")

        # Eksekusi fungsi ekstrak_bagian menggunakan dokumen yang sudah dibaca
        hasil = ekstrak_bagian(doc)

        # Jika hasil adalah string, berarti ada bagian yang kosong
        if isinstance(hasil, str):
            print(hasil)
        else:
            # Daftar bagian yang akan diuji
            bagian_keys = [
                "Abstrak",
                "Abstract",
                "Pendahuluan",
                "Metode Penelitian",
                "Hasil dan Pembahasan",
                "Kesimpulan",
                "Referensi"
            ]

            # Periksa setiap bagian, pastikan tidak kosong
            for key in bagian_keys:
                with self.subTest(bagian=key):
                    self.assertTrue(hasil[key], f"Bagian {key} seharusnya tidak kosong.")

    def test_generate_with_template(self):
        logging.info("Memulai test_generate_dengan_template_benar...")
        # Baca template dokumen
        doc_template = baca_template("template_jurnal/Template_test.docx")
        bagian = ekstrak_bagian(Document("data/test/sample_test.docx"))

        # Jalankan fungsi untuk menyesuaikan dengan template
        result = sesuaikan_dengan_template(doc_template, bagian)
        
        # Simpan hasil ke file sementara untuk verifikasi
        temp_output_path = 'temp_test_output.docx'
        doc_template.save(temp_output_path)
        
        # Pastikan dokumen berhasil di-generate tanpa error
        self.assertNotEqual(result, 'Template tidak sesuai')
        
        # Periksa apakah file output berhasil di-generate
        self.assertTrue(os.path.exists(temp_output_path))
        
        # Hapus file sementara setelah test
        os.remove(temp_output_path)

        logging.info("test_generate_dengan_template_yang_benar berhasil.\n")
        
if __name__ == "__main__":
    unittest.main()
